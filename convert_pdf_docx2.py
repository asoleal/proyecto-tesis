import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def pdf_a_docx(pdf_path, docx_path):
    # Crear un nuevo documento de Word
    doc = Document()

    # Abrir el archivo PDF
    pdf_document = fitz.open(pdf_path)

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)  # Cargar la página
        text_instances = page.get_text("blocks")  # Extraer bloques de texto

        # Agregar texto al documento Word
        for block in sorted(text_instances, key=lambda x: (x[1], x[0])):  # Ordenar por posición
            text = block[4].strip()  # El texto está en el índice 4
            if text:
                doc.add_paragraph(text)

        # Extraer imágenes de la página
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]  # Referencia de la imagen
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]

            # Guardar temporalmente la imagen
            image_path = f"temp_image_{page_num}_{img_index}.png"
            with open(image_path, "wb") as image_file:
                image_file.write(image_bytes)

            # Agregar la imagen al documento Word
            doc.add_picture(image_path, width=Inches(4))  # Ajusta el ancho según sea necesario

    # Guardar el documento Word
    doc.save(docx_path)
    print(f"El archivo DOCX ha sido guardado en: {docx_path}")

# Rutas de los archivos
pdf_path = "0000.pdf"  # Cambia esto por la ruta de tu archivo PDF
docx_path = "salida.docx"  # Cambia esto por la ruta donde deseas guardar el archivo DOCX

# Llamar a la función para convertir
pdf_a_docx(pdf_path, docx_path)
